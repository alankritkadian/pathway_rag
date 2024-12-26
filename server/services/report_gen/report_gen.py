import os
import yfinance as yf
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from openai import OpenAI
import re
import numpy as np
# import comtypes.client
from pydantic import BaseModel, Field#type: ignore
from langchain.tools import StructuredTool

# Load OpenAI API key and Finnhub API key from environment variables
os.environ["OPENAI_API_KEY"] = "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"
os.environ["FINNHUB_API_KEY"] = "ct8o3k1r01qtkv5spi7gct8o3k1r01qtkv5spi80"  # Replace with your Finnhub API key

# Initialize the OpenAI client
client = OpenAI()

# def convert_docx_to_pdf(input_path: str, output_path: str) -> None:
#     """
#     Converts a .docx file to .pdf format using Microsoft Word via comtypes.

#     Args:
#         input_path (str): Path to the input .docx file.
#         output_path (str): Path to save the output .pdf file.

#     Returns:
#         None
#     """
#     try:
#         # Initialize Word application
#         word = comtypes.client.CreateObject("Word.Application")
#         word.Visible = False

#         # Open the Word document
#         doc = word.Documents.Open(input_path)

#         # Save as PDF
#         doc.SaveAs(output_path, FileFormat=17)  # 17 corresponds to PDF format in Word

#         # Close the document and quit Word
#         doc.Close()
#         word.Quit()

#         print(f"Conversion successful! PDF saved at: {output_path}")
#     except Exception as e:
#         print(f"An error occurred: {e}")



def chat_client(query: str) -> str:
    """
    A simple chat client that uses OpenAI's GPT model to answer a query.

    :param query: The user's question or prompt.
    :return: The AI's response as a string.
    """
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": query}]
    )
    return response.choices[0].message.content.strip()

def extract_ticker_name(query):
    """
    Extract company name from the user query using LLM.

    :return: Extracted company name
    """
    try:
        prompt = f"""Extract the stock ticker for the company mentioned in the query. Use the examples below as context:
          Apple Inc. -> AAPL
          Microsoft Corporation -> MSFT
          Tesla, Inc. -> TSLA
          Amazon.com, Inc. -> AMZN
          Alphabet Inc. (Class A) -> GOOGL
          Meta Platforms, Inc. -> META
          Nvidia Corporation -> NVDA
          Berkshire Hathaway Inc. -> BRK.B
          Johnson & Johnson -> JNJ
          Visa Inc. -> V
          Procter & Gamble Co. -> PG
          JPMorgan Chase & Co. -> JPM
          Exxon Mobil Corporation -> XOM
          Samsung Electronics -> SSNLF
          Toyota Motor Corporation -> TM
          Walmart Inc. -> WMT
          Pfizer Inc. -> PFE
          Coca-Cola Company -> KO
          Intel Corporation -> INTC
          AT&T Inc. -> T
          Verizon Communications Inc. -> VZ
          Bank of America Corporation -> BAC
          Mastercard Incorporated -> MA
          Netflix, Inc. -> NFLX
          Adobe Inc. -> ADBE
          PayPal Holdings, Inc. -> PYPL
          Cisco Systems, Inc. -> CSCO
          Oracle Corporation -> ORCL
          IBM Corporation -> IBM
          General Electric Company -> GE
          Chevron Corporation -> CVX
          Morgan Stanley -> MS
          Goldman Sachs Group, Inc. -> GS
          UnitedHealth Group Incorporated -> UNH
          Home Depot, Inc. -> HD
          McDonald's Corporation -> MCD
          PepsiCo, Inc. -> PEP
          Comcast Corporation -> CMCSA
          Disney (The Walt Disney Company) -> DIS
          Salesforce, Inc. -> CRM
          Twitter, Inc. -> TWTR
          Uber Technologies, Inc. -> UBER
          Airbnb, Inc. -> ABNB
          Spotify Technology S.A. -> SPOT
          Shopify Inc. -> SHOP
          Square, Inc. -> SQ
          Snap Inc. -> SNAP
          Palantir Technologies Inc. -> PLTR
          Snowflake Inc. -> SNOW
          Moderna, Inc. -> MRNA
          Zoom Video Communications, Inc. -> ZM
          DocuSign, Inc. -> DOCU
          Okta, Inc. -> OKTA
          Twilio Inc. -> TWLO
          Slack Technologies, Inc. -> WORK
          Pinterest, Inc. -> PINS
          Roku, Inc. -> ROKU
          Peloton Interactive, Inc. -> PTON
          Fiverr International Ltd. -> FVRR
          Etsy, Inc. -> ETSY
          Zillow Group, Inc. -> ZG
          Wayfair Inc. -> W
          Beyond Meat, Inc. -> BYND
          DoorDash, Inc. -> DASH
          Roblox Corporation -> RBLX
          Unity Software Inc. -> U
          Palantir Technologies Inc. -> PLTR
          Coinbase Global, Inc. -> COIN
          DraftKings Inc. -> DKNG
          Fisker Inc. -> FSR
          Lucid Group, Inc. -> LCID
          Rivian Automotive, Inc. -> RIVN
          Nikola Corporation -> NKLA
          Virgin Galactic Holdings, Inc. -> SPCE
          NIO Inc. -> NIO
          XPeng Inc. -> XPEV
          Li Auto Inc. -> LI
          BYD Company Ltd. -> BYDDY
          Tencent Holdings Ltd. -> TCEHY
          Alibaba Group Holding Ltd. -> BABA
          Baidu, Inc. -> BIDU
          JD.com, Inc. -> JD
          Meituan -> MPNGF
          Xiaomi Corporation -> XIACY
          Infosys Ltd. -> INFY
          Tata Consultancy Services Ltd. -> TCS
          HDFC Bank Ltd. -> HDB
          Reliance Industries Ltd. -> RELIANCE
          ICICI Bank Ltd. -> IBN
          State Bank of India -> SBIN
          just the ticker no other symbol like fullstop: '{query}'"""
        print(f"Prompt sent to LLM: {prompt}")  # Debug print
        company_name = chat_client(prompt)
        print(f"Extracted company name: {company_name}")  # Debug print
        return company_name
    except Exception as e:
        raise ValueError(f"Error extracting company name: {str(e)}")

def process_line(line):
    """
    Process a line for additional formatting:
    - Convert # headers to bold text.
    - Make text enclosed in ** bold and remove **.
    - Remove -- and replace with an empty line.
    """
    line = line.strip()

    # Remove double dashes and replace with empty line
    if line == '--':
        return None

    # Handle headers (lines starting with #)
    if line.startswith('#'):
        return {'type': 'header', 'content': line.lstrip('#').strip()}

    # Handle bold text within **
    if '*' in line:
        bold_segments = []
        parts = line.split('*')
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Odd parts are within **
                bold_segments.append({'type': 'bold', 'content': part})
            else:
                bold_segments.append({'type': 'normal', 'content': part})
        return bold_segments

    # Return the line as a normal paragraph
    return {'type': 'normal', 'content': line}

class EnhancedFinancialDataAnalyzer:
    def __init__(self, ticker_symbol):
        self.ticker_symbol = ticker_symbol
        self.stock = yf.Ticker(ticker_symbol)
        self.hist_data = self.stock.history(period="10y")  # Fetch 10 years of data
        self.hist_data.index = pd.to_datetime(self.hist_data.index)

        # Load OpenAI API key from environment variable
        load_dotenv()
        self.client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    def plot_stock_price(self):
        plt.figure(figsize=(12, 6))
        plt.plot(self.hist_data.index, self.hist_data['Close'], label='Close Price', color='blue')
        ma_50 = self.hist_data['Close'].rolling(window=50).mean()
        ma_200 = self.hist_data['Close'].rolling(window=200).mean()
        plt.plot(ma_50.index, ma_50, label='50-day MA', color='red', linewidth=2)
        plt.plot(ma_200.index, ma_200, label='200-day MA', color='green', linewidth=2)
        plt.title(f"{self.ticker_symbol} Stock Price with Moving Averages")
        plt.xlabel('Date')
        plt.ylabel('Price')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_stock_price.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_daily_returns(self):
        daily_returns = self.hist_data['Close'].pct_change()
        plt.figure(figsize=(12, 6))
        plt.plot(daily_returns.index, daily_returns * 100, label='Daily Returns', color='purple')
        plt.axhline(0, color='gray', linestyle='--')
        plt.title(f"{self.ticker_symbol} Daily Returns")
        plt.xlabel('Date')
        plt.ylabel('Daily Return (%)')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_daily_returns.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_bollinger_bands(self):
        window = 20
        std_dev = self.hist_data['Close'].rolling(window=window).std()
        middle_band = self.hist_data['Close'].rolling(window=window).mean()
        upper_band = middle_band + (std_dev * 2)
        lower_band = middle_band - (std_dev * 2)

        plt.figure(figsize=(12, 6))
        plt.plot(self.hist_data.index, self.hist_data['Close'], label='Close Price', color='blue')
        plt.plot(middle_band.index, middle_band, label='Middle Band', color='gray', linestyle='--')
        plt.plot(upper_band.index, upper_band, label='Upper Band', color='red', linewidth=1)
        plt.plot(lower_band.index, lower_band, label='Lower Band', color='green', linewidth=1)
        plt.fill_between(lower_band.index, lower_band, upper_band, color='green', alpha=0.2)
        plt.title(f"{self.ticker_symbol} Bollinger Bands")
        plt.xlabel('Date')
        plt.ylabel('Price')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_bollinger_bands.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_volume_analysis(self):
        plt.figure(figsize=(12, 6))
        plt.bar(self.hist_data.index, self.hist_data['Volume'], color='blue', alpha=0.7)
        plt.title(f"{self.ticker_symbol} Trading Volume")
        plt.xlabel('Date')
        plt.ylabel('Volume')
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_volume_analysis.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_rsi(self):
        delta = self.hist_data['Close'].diff()
        gain = (delta.where(delta > 0, 0)).rolling(window=14).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(window=14).mean()
        rs = gain / loss
        rsi = 100 - (100 / (1 + rs))

        plt.figure(figsize=(12, 6))
        plt.plot(rsi.index, rsi, label='RSI', color='purple')
        plt.axhline(30, color='red', linestyle='--', label='Oversold')
        plt.axhline(70, color='green', linestyle='--', label='Overbought')
        plt.title(f"{self.ticker_symbol} RSI (Relative Strength Index)")
        plt.xlabel('Date')
        plt.ylabel('RSI')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_rsi.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_macd(self):
        exp1 = self.hist_data['Close'].ewm(span=12, adjust=False).mean()
        exp2 = self.hist_data['Close'].ewm(span=26, adjust=False).mean()
        macd = exp1 - exp2
        signal = macd.ewm(span=9, adjust=False).mean()

        plt.figure(figsize=(12, 6))
        plt.plot(macd.index, macd, label='MACD', color='blue')
        plt.plot(signal.index, signal, label='Signal Line', color='red')
        plt.title(f"{self.ticker_symbol} MACD (Moving Average Convergence Divergence)")
        plt.xlabel('Date')
        plt.ylabel('MACD')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_macd.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_stochastic_oscillator(self):
        low_14 = self.hist_data['Low'].rolling(window=14).min()
        high_14 = self.hist_data['High'].rolling(window=14).max()
        k_percent = 100 * (self.hist_data['Close'] - low_14) / (high_14 - low_14)
        d_percent = k_percent.rolling(window=3).mean()

        plt.figure(figsize=(12, 6))
        plt.plot(k_percent.index, k_percent, label='%K', color='blue')
        plt.plot(d_percent.index, d_percent, label='%D', color='red')
        plt.axhline(20, color='green', linestyle='--', label='Oversold')
        plt.axhline(80, color='red', linestyle='--', label='Overbought')
        plt.title(f"{self.ticker_symbol} Stochastic Oscillator")
        plt.xlabel('Date')
        plt.ylabel('Stochastic %')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_stochastic_oscillator.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_obv(self):
        obv = np.where(self.hist_data['Close'] > self.hist_data['Close'].shift(1),
                       self.hist_data['Volume'],
                       np.where(self.hist_data['Close'] < self.hist_data['Close'].shift(1),
                                -self.hist_data['Volume'], 0)).cumsum()

        plt.figure(figsize=(12, 6))
        plt.plot(self.hist_data.index, obv, label='OBV', color='blue')
        plt.title(f"{self.ticker_symbol} On-Balance Volume (OBV)")
        plt.xlabel('Date')
        plt.ylabel('OBV')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_obv.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def plot_atr(self):
        high_low = self.hist_data['High'] - self.hist_data['Low']
        high_close = np.abs(self.hist_data['High'] - self.hist_data['Close'].shift())
        low_close = np.abs(self.hist_data['Low'] - self.hist_data['Close'].shift())
        ranges = pd.concat([high_low, high_close, low_close], axis=1)
        true_range = ranges.max(axis=1)
        atr = true_range.rolling(window=14).mean()

        plt.figure(figsize=(12, 6))
        plt.plot(atr.index, atr, label='ATR', color='blue')
        plt.title(f"{self.ticker_symbol} Average True Range (ATR)")
        plt.xlabel('Date')
        plt.ylabel('ATR')
        plt.legend()
        plt.tight_layout()
        fig_path = f"{self.ticker_symbol}_atr.png"
        plt.savefig(fig_path)
        plt.close()
        return fig_path

    def generate_llm_analysis(self, company_info, query):
        """
        Generate comprehensive financial analysis using GPT-4
        """
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a financial analyst providing detailed company reports."},
                    {"role": "user", "content": self.generate_comprehensive_prompt(company_info, query)}
                ],
                max_tokens=4000,
                temperature=0.7
            )
            analysis_text = response.choices[0].message.content
            print("LLM Analysis Generated: ")
            print(analysis_text)
            return analysis_text
        except Exception as e:
            return f"Error generating analysis: {str(e)}"
    def generate_comprehensive_prompt(self, company_info, query):
        """
        Enhanced prompt for more targeted and real-time financial analysis
        """
        current_price = company_info.get('currentPrice', 'N/A')
        market_cap = company_info.get('marketCap', 'N/A')
        pe_ratio = company_info.get('trailingPE', 'N/A')
        dividend_yield = company_info.get('dividendYield', 'N/A')

        # Handle dividend_yield formatting
        if dividend_yield != 'N/A':
            try:
                dividend_yield = float(dividend_yield)
                dividend_yield_str = f"{dividend_yield * 100:.2f}%"
            except ValueError:
                dividend_yield_str = 'N/A'
        else:
            dividend_yield_str = 'N/A'

        # return f"""
        # Please follow the specific instructions if it mentions, if it does not mention then follow this template in detail. Please note that the content must be more than one page.

        # Provide a comprehensive, concise financial analysis for {company_info.get('longName', self.ticker_symbol)} with real-time market context and actionable insights:

        # Current Market Snapshot
        # - Current Price: ${current_price}
        # - Market Cap: ${market_cap:,}
        # - P/E Ratio: {pe_ratio}
        # - Dividend Yield: {dividend_yield_str}

        # QUERY CLASSIFICATIONS:
        # - Company Financial Overview: Analyze current metrics, historical performance, market position, and future outlook.
        # - Comparative Financial Analysis: Compare financials with competitors, assessing market strength and growth.
        # - Historical Performance: Show trends in financial metrics, key KPIs, and market impact over time.
        # - Investment Assessment: Evaluate current valuation, growth prospects, risks, and analyst recommendations.
        # - Specific Metric Queries: Provide detailed analysis of metrics like P/E ratio, market cap, etc.
        # - Sector/Industry Analysis: Analyze sector trends, key players, and industry health.
        # - Predictive Scenarios: Provide financial forecasts based on trends and potential future market changes.

        # Please keep the above points in paragraph format and focus more on the user query, with the response tailored based on the user request.
        # """
        
        return f"""
        Ensure the response is tailored to the query, addressing it comprehensively and maintaining a professional format.
        For all queries, regardless of complexity, generate a comprehensive, two and a half (2.5) page report with well-defined sections and subsections to ensure clarity and logical flow. Each report should be tailored to the query, emphasizing the specifics of what the user has asked for, and providing an in-depth, professional response. Even for straightforward queries like "generate the financial report," ensure the report includes detailed analysis, covering all relevant aspects in an organized format.  
        We want to satisfy the user by generating content that satisfies his query.
        User query: {query}
        Provide a focused and detailed financial analysis for {company_info.get('longName', self.ticker_symbol)} with actionable insights based on real-time market context. The response should prioritize the user's query, ensuring that all requested details are addressed thoroughly. 
        Make a 3-4 liner text for the part mentioned under key financial highlight, not more than that. Focus more on user query and generate relevant response to what user have asked in the question.
        Key financial highlights include:
        - Current Price: ${current_price}
        - Market Cap: ${market_cap:,}
        - P/E Ratio: {pe_ratio}
        - Dividend Yield: {dividend_yield_str} 
        
       For all queries, regardless of complexity, generate a comprehensive, two and a half (2.5) page report with well-defined sections and subsections to ensure clarity and logical flow. Each report should be tailored to the query, emphasizing the specifics of what the user has asked for, and providing an in-depth, professional response. Even for straightforward queries like "generate the financial report," ensure the report includes detailed analysis, covering all relevant aspects in an organized format.

        # Please keep the above points in paragraph format and focus more on the user query, with the response tailored based on the user request.
        
        """
    def generate_comprehensive_report(self, query):
        # Create a temporary folder
        temp_folder = "Financial_Reports"
        os.makedirs(temp_folder, exist_ok=True)

        # Fetch company information
        company_info = self.stock.info

        # Generate financial visualizations
        visualizations = {}
        plot_methods = [
            self.plot_stock_price,
            self.plot_daily_returns,
            self.plot_bollinger_bands,
            self.plot_volume_analysis,
            self.plot_rsi,
            self.plot_macd,
            self.plot_stochastic_oscillator,
            self.plot_obv,
            self.plot_atr
        ]

        for plot_method in plot_methods:
            try:
                result = plot_method()
                if result:
                    visualizations[plot_method.__name__.replace('plot', '')] = result
            except Exception as e:
                print(f"Plot generation error: {e}")

        # Generate LLM analysis
        llm_analysis = self.generate_llm_analysis(company_info, query)
        print(llm_analysis)

        # Generate report
        self.create_comprehensive_docx_report(company_info, llm_analysis, visualizations, temp_folder)

        # Clean up: Delete images and DOCX
        for file_path in visualizations.values():
            os.remove(file_path)

        print("Report generated successfully.")

    def create_comprehensive_docx_report(self, company_info, llm_analysis, visualizations, temp_folder):
        doc = Document()

        # Set minimal margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(36)  # 0.5 inch
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(36)
            section.right_margin = Pt(36)

        # Title and Header
        title = doc.add_heading(f'Financial Analysis: {company_info.get("longName", self.ticker_symbol)}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add basic company info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        info_para.add_run(f'Ticker: {self.ticker_symbol} | ').bold = True
        info_para.add_run(f'Sector: {company_info.get("sector", "N/A")} | ').bold = True
        info_para.add_run(f'Market Cap: ${company_info.get("marketCap", 0):,}').bold = True

        # Create a two-column section
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Left column for text (2/3 width)
        left_cell = table.cell(0, 0)
        left_cell.width = Inches(6)  # Approximately 2/3 of page width

        # Right column for images (1/3 width)
        right_cell = table.cell(0, 1)
        right_cell.width = Inches(3)  # Approximately 1/3 of page width

        # Set default font and size
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Add LLM analysis to left cell
        self._add_formatted_text(left_cell, llm_analysis)

        # Add visualizations to right cell
        for key, file_path in visualizations.items():
            try:
                para = right_cell.add_paragraph()  # Add a new paragraph to the cell
                para.add_run(f"\n{key.replace('_', ' ').title()}:").bold = True
                para.add_run().add_picture(file_path, width=Inches(2.5))  # Add picture to the paragraph
                para.add_run("\n\n")  # Add new lines or spaces for manual checking
            except Exception as e:
                print(f"Visualization error: {e}")

        # Save the document in the temporary folder
        docx_path = os.path.join(temp_folder, f"{self.ticker_symbol}_financial_report.docx")
        doc.save(docx_path)
        print("&&&&&&&&&&&&&&&&&&&&")
        docx_path=os.getcwd()+"/"+docx_path
        print(docx_path)
        # convert_docx_to_pdf(docx_path, docx_path.replace('.docx', '.pdf'))

    def _add_formatted_text(self, cell, text):
        """
        Add formatted text to a Word document cell based on specific formatting rules.
        """
        # Split the text into lines
        lines = text.split('\n')
        for line in lines:
            # Process each line with the temp function
            processed_line = process_line(line)
            if not processed_line:
                continue

            # Handle processed output
            if isinstance(processed_line, dict):  # Single processed line
                if processed_line['type'] == 'header':
                    para = cell.add_paragraph()
                    run = para.add_run(processed_line['content'])
                    run.bold = True
                    run.font.size = Pt(14)  # Optional: Adjust font size for headers
                elif processed_line['type'] == 'normal':
                    cell.add_paragraph(processed_line['content'])
                elif processed_line['type'] == 'bold':
                    para = cell.add_paragraph()
                    run = para.add_run(processed_line['content'])
                    run.bold = True
            elif isinstance(processed_line, list):  # Mixed normal and bold segments
                para = cell.add_paragraph()
                for segment in processed_line:
                    run = para.add_run(segment['content'])
                    if segment['type'] == 'bold':
                        run.bold = True

# Usage example
if __name__ == "__main__":
    query = input("Enter your query (e.g., 'Generate a report for Tesla'): ")
    ticker = extract_ticker_name(query)
    analyzer = EnhancedFinancialDataAnalyzer(ticker)
    analyzer.generate_comprehensive_report(query)

class ReporGenNode(BaseModel):
    query: str = Field(description="The Query containing companies name to be processed for Report Generation")


def report_node_function(query: str) -> str:
    """
    An LLM agent with access to a structured tool that generates comprehensive reports on query provided 
    which contains the companies name and generates report from data fetched form its internal tools. 
    """
    try:
        # Step 1: Extract the ticker from the query
        ticker = extract_ticker_name(query)
        if not ticker:
            return "ERROR: Unable to extract a ticker from the query. Please provide a valid input."
    
        print(f"Ticker '{ticker}' extracted successfully.")
    
        # Step 2: Initialize the analyzer and generate the report
        analyzer = EnhancedFinancialDataAnalyzer(ticker)
        analyzer.generate_comprehensive_report(query)
    
        return "SUCCESS: The financial report has been generated successfully."

    except ValueError as ve:
        # Handle specific input-related errors
        return f"ERROR: Invalid input or query - {ve}"

    except KeyError as ke:
        # Handle missing data for the given ticker
        return f"ERROR: Data for ticker '{ticker}' is incomplete or unavailable - {ke}"

    except Exception as e:
        # Handle any unexpected exceptions
        error_message = f"ERROR: An unexpected error occurred while generating the report - {e}"
        # Consider logging the full traceback for debugging purposes
        print(error_message)
        return error_message



# reportgen_tool = StructuredTool.from_function(
#     report_node_function,
#     name="reportgen_tool",
#     description="""reportgen_tool(query: str) -> str:
#     NEVER use this tool unless you are asked to generate a report explicitly.
#     ALWAYS REMEMBER: this tool can only be used to generate reports for the user not for information retrieval, for information retrieval refer to the data node tool.
#     This tool cannot access specific data to answer a user's query, it just fetches data from certain apis to generate reports, so it should be used only when the user explicitly asks for a report.
#     Only that part of query should be passed to this tool which demands for report generation. Other forms of data requested should be handled by other tools.
#      An LLM agent with access to a structured tool that generates comprehensive reports on query provided 
#     which contains the companies name, user query and generates report from data fetched form its internal tools. 
#     Provide concise queries to this tool which contains company name AND always include the user query details, DO NOT give vague queries for search like
#     - 'Generate report for the company whose CEO is Elon Musk'
#     - 'Generate report for the most valued company of the world'
#     Instead, provide specific queries with what exactly the user demanded for like
#     - 'Report for Tesla'
#     - 'Report for Tesla for financial year 2021'
#     - 'Report for Tesla and analysis of their autonomous vechile division'
#     ALWAYS mention company name for generating reports
#     ALWAYS provide with the user query details to get the accurate results (in case user query wasn't too specific, don't add any aditional details)
#     PROVIDE with what exactly the user has demanded for in the query
#     Eg: Always use standard names for companies while sending queries.
#     ALWAYS provide specific queries to get accurate results.
#     ENSURE to keep in mind, the number of reports the user wants. Don't overcall this tool ever (if the user wants one report, 
#     this tool shall be call only one. Depending upon the company names or how many reports specfically wants - this tool shall be called)
#     DO NOT try to fetch multiple results in a single query, instead, make multiple queries.
#     NEVER use this tool unless you are asked to generate a report explicitly.
#     """,
#     args_schema=ReporGenNode,
# )

